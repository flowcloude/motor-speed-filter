import docx

doc = docx.Document('速度估计与滤波方法_课设报告_修改版.docx')

print("=== 修改后的关键段落 ===")
for i, para in enumerate(doc.paragraphs):
    if 'Motor Pilot' in para.text or '528' in para.text:
        print(f"段落 {i}: {para.text[:150]}")

print("\n=== 表格中是否还有522 ===")
for i, table in enumerate(doc.tables):
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        if '522' in str(cells):
            print(f"表格 {i} 包含 522: {cells}")
print("检查完成")
