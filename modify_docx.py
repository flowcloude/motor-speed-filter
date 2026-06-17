import docx

doc = docx.Document('速度估计与滤波方法_课设报告.docx')

replacements = [
    ('手动 Watch 窗口数据采集', 'Motor Pilot 实时波形显示'),
    ('Keil Watch 窗口查看', 'Motor Pilot 波形窗口实时观察'),
    ('Watch窗口', 'Motor Pilot'),
]

for i, para in enumerate(doc.paragraphs):
    for old, new in replacements:
        if old in para.text:
            para.text = para.text.replace(old, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == '522':
                cell.text = '528'

doc.save('速度估计与滤波方法_课设报告_修改版.docx')
print("文档修改完成！")
