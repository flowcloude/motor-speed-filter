import docx

doc = docx.Document('速度估计与滤波方法_课设报告.docx')

replacements = [
    ('522 RPM', '528 RPM'),
    ('约 522 RPM', '约 528 RPM'),
    ('手动 Watch 窗口数据采集', 'Motor Pilot 实时波形显示'),
    ('Keil Debug Watch 窗口导出', 'Motor Pilot 波形窗口实时观察'),
    ('Keil Watch 窗口查看', 'Motor Pilot 波形窗口实时观察'),
    ('Watch窗口', 'Motor Pilot'),
]

for para in doc.paragraphs:
    for old, new in replacements:
        if old in para.text:
            para.text = para.text.replace(old, new)

for para in doc.paragraphs:
    if '当 raw=88 达到稳态时' in para.text:
        para.text = para.text.replace('当 raw=88 达到稳态时，LPF 输出仅 86', '当原始速度约为 530 RPM 达到稳态时，LPF 输出约为 527 RPM')
    if 'raw 92→64' in para.text:
        para.text = para.text.replace('raw 92→64', '原始速度阶跃变化')

doc.save('速度估计与滤波方法_课设报告_修改版.docx')
print("文档修改完成！")
