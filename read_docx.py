import docx

doc = docx.Document('速度估计与滤波方法_课设报告.docx')

for i, para in enumerate(doc.paragraphs):
    print(f"段落 {i}: {para.text[:200]}")

print("\n\n=== 表格内容 ===")
for i, table in enumerate(doc.tables):
    print(f"\n表格 {i}:")
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        print(f"  {cells}")
